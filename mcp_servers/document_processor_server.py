#!/usr/bin/env python3
"""
MCP Document Processor Server — demonstrates document parsing, content extraction,
format detection, text analysis, and multi-format document handling.

Install: pip install mcp PyPDF2 python-docx beautifulsoup4 markdown
Configure in .claude/settings.json:
{
  "mcpServers": {
    "docs": {
      "type": "stdio",
      "command": "python",
      "args": ["mcp_servers/document_processor_server.py"]
    }
  }
}
"""

import asyncio
import csv
import io
import json
import os
import re
from collections import Counter
from pathlib import Path
from typing import Any

from mcp.server import Server
from mcp.server.stdio import StdioServerTransport
from mcp.types import Tool, TextContent

server = Server("docs")

# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _read_text_file(path: str) -> str:
    p = Path(path)
    if not p.exists():
        raise FileNotFoundError(f"File not found: {path}")
    return p.read_text(encoding="utf-8", errors="replace")

def _extract_text_from_pdf(file_bytes: bytes) -> str:
    try:
        from PyPDF2 import PdfReader
        reader = PdfReader(io.BytesIO(file_bytes))
        pages = []
        for page in reader.pages:
            text = page.extract_text()
            if text:
                pages.append(text)
        return "\n\n".join(pages)
    except ImportError:
        return "[PyPDF2 not installed. Run: pip install PyPDF2]"

def _extract_text_from_docx(file_bytes: bytes) -> str:
    try:
        from docx import Document
        doc = Document(io.BytesIO(file_bytes))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        return "\n".join(paragraphs)
    except ImportError:
        return "[python-docx not installed. Run: pip install python-docx]"

def _extract_text_from_html(html: str) -> str:
    try:
        from bs4 import BeautifulSoup
        soup = BeautifulSoup(html, "html.parser")
        for tag in soup(["script", "style", "nav", "footer", "header"]):
            tag.decompose()
        return soup.get_text(separator="\n", strip=True)
    except ImportError:
        # Fallback regex-based extraction
        text = re.sub(r"<script[^>]*>.*?</script>", "", html, flags=re.DOTALL | re.IGNORECASE)
        text = re.sub(r"<style[^>]*>.*?</style>", "", text, flags=re.DOTALL | re.IGNORECASE)
        text = re.sub(r"<[^>]+>", " ", text)
        text = re.sub(r"\s+", " ", text).strip()
        return text

def _detect_format(filename: str) -> str:
    ext = Path(filename).suffix.lower()
    formats = {
        ".txt": "text",
        ".md": "markdown",
        ".json": "json",
        ".csv": "csv",
        ".xml": "xml",
        ".html": "html",
        ".htm": "html",
        ".pdf": "pdf",
        ".docx": "docx",
        ".doc": "doc",
        ".py": "python",
        ".js": "javascript",
        ".ts": "typescript",
        ".rs": "rust",
        ".go": "go",
        ".java": "java",
        ".yaml": "yaml",
        ".yml": "yaml",
        ".toml": "toml",
        ".ini": "ini",
        ".cfg": "ini",
        ".log": "log",
    }
    return formats.get(ext, "unknown")

# ---------------------------------------------------------------------------
# Tools
# ---------------------------------------------------------------------------

@server.tool()
async def parse_document(
    file_path: str = "",
    content: str = "",
    format: str = "auto",
) -> list[TextContent]:
    """Parse and extract text content from a document file.

    Args:
        file_path: Path to the document file
        content: Raw document content as string (used if file_path is empty)
        format: "auto" (detect from extension), "txt", "md", "html", "pdf", "docx", "json", "csv"
    """
    if file_path:
        fmt = _detect_format(file_path)
        if format != "auto":
            fmt = format

        if fmt == "pdf":
            text = _extract_text_from_pdf(Path(file_path).read_bytes())
        elif fmt == "docx":
            text = _extract_text_from_docx(Path(file_path).read_bytes())
        elif fmt == "html":
            text = _extract_text_from_html(Path(file_path).read_text(errors="replace"))
        else:
            text = Path(file_path).read_text(errors="replace")
    elif content:
        fmt = format if format != "auto" else "text"
        if fmt == "html":
            text = _extract_text_from_html(content)
        else:
            text = content
    else:
        return [TextContent(type="text",
            text=json.dumps({"error": "Provide either file_path or content"}, indent=2))]

    # Basic stats
    words = text.split()
    lines = text.splitlines()
    chars = len(text)

    return [TextContent(type="text", text=json.dumps({
        "format": fmt,
        "file": file_path or "<inline>",
        "stats": {
            "characters": chars,
            "words": len(words),
            "lines": len(lines),
            "paragraphs": len([l for l in lines if not l.strip()]) + 1,
            "estimated_reading_time_minutes": max(1, len(words) // 200),
        },
        "text": text[:10000] + (f"\n\n[... {chars - 10000} more characters]" if chars > 10000 else ""),
    }, indent=2, ensure_ascii=False))]

@server.tool()
async def extract_metadata(file_path: str) -> list[TextContent]:
    """Extract metadata from a document file (author, dates, title, etc.).

    Args:
        file_path: Path to the document
    """
    p = Path(file_path)
    if not p.exists():
        return [TextContent(type="text",
            text=json.dumps({"error": f"File not found: {file_path}"}, indent=2))]

    fmt = _detect_format(file_path)
    meta: dict[str, Any] = {
        "filename": p.name,
        "extension": p.suffix,
        "size_bytes": p.stat().st_size,
        "size_human": f"{p.stat().st_size / 1024:.1f} KB" if p.stat().st_size < 1024**2
                      else f"{p.stat().st_size / (1024**2):.1f} MB",
        "created": datetime.fromtimestamp(p.stat().st_ctime).isoformat(),
        "modified": datetime.fromtimestamp(p.stat().st_mtime).isoformat(),
        "detected_format": fmt,
    }

    if fmt == "pdf":
        try:
            from PyPDF2 import PdfReader
            reader = PdfReader(str(p))
            if reader.metadata:
                meta["pdf_metadata"] = {
                    "title": reader.metadata.get("/Title", ""),
                    "author": reader.metadata.get("/Author", ""),
                    "subject": reader.metadata.get("/Subject", ""),
                    "creator": reader.metadata.get("/Creator", ""),
                    "pages": len(reader.pages),
                }
        except ImportError:
            meta["pdf_metadata"] = "PyPDF2 not installed"
        except Exception as e:
            meta["pdf_error"] = str(e)

    elif fmt == "docx":
        try:
            from docx import Document
            doc = Document(str(p))
            core = doc.core_properties
            meta["docx_metadata"] = {
                "title": core.title or "",
                "author": core.author or "",
                "created": str(core.created) if core.created else "",
                "modified": str(core.modified) if core.modified else "",
                "paragraphs": len(doc.paragraphs),
            }
        except ImportError:
            meta["docx_metadata"] = "python-docx not installed"
        except Exception as e:
            meta["docx_error"] = str(e)

    return [TextContent(type="text", text=json.dumps(meta, indent=2, default=str))]

@server.tool()
async def text_analysis(
    text: str,
    analysis_type: str = "all",
) -> list[TextContent]:
    """Analyze text: word frequency, readability, sentiment, entity extraction.

    Args:
        text: The text to analyze
        analysis_type: "all", "frequency", "readability", "entities"
    """
    from datetime import datetime

    result: dict[str, Any] = {"text_length": len(text), "word_count": len(text.split())}
    words = re.findall(r"\b[a-zA-Z]+\b", text.lower())

    if analysis_type in ("all", "frequency"):
        # Word frequency
        word_freq = Counter(words).most_common(50)
        # Bigrams
        bigrams = Counter(zip(words, words[1:])).most_common(20)
        result["frequency"] = {
            "top_words": [{"word": w, "count": c} for w, c in word_freq if len(w) > 3],
            "top_bigrams": [{"bigram": f"{a} {b}", "count": c} for (a, b), c in bigrams],
            "unique_words": len(set(words)),
            "lexical_diversity": round(len(set(words)) / max(len(words), 1), 3),
        }

    if analysis_type in ("all", "readability"):
        sentences = re.split(r"[.!?]+", text)
        sentences = [s.strip() for s in sentences if s.strip()]
        if sentences and words:
            avg_words_per_sentence = len(words) / len(sentences)
            # Flesch-Kincaid Grade Level approximation
            total_syllables = sum(_count_syllables(w) for w in words)
            if len(sentences) > 0:
                fk_grade = 0.39 * (len(words) / len(sentences)) + 11.8 * (total_syllables / len(words)) - 15.59
            else:
                fk_grade = 0
            result["readability"] = {
                "sentence_count": len(sentences),
                "avg_sentence_length": round(avg_words_per_sentence, 1),
                "flesch_kincaid_grade": round(max(0, fk_grade), 1),
                "flesch_reading_ease": round(206.835 - 1.015 * (len(words) / len(sentences))
                                             - 84.6 * (total_syllables / len(words)), 1) if sentences else 0,
            }

    if analysis_type in ("all", "entities"):
        # Simple entity extraction
        emails = list(set(re.findall(r"[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}", text)))
        urls = list(set(re.findall(r"https?://[^\s]+", text)))
        # Potential names (capitalized word pairs)
        names = list(set(re.findall(r"\b[A-Z][a-z]+\s[A-Z][a-z]+\b", text)))[:20]
        dates = list(set(re.findall(r"\d{4}-\d{2}-\d{2}|\d{2}/\d{2}/\d{4}", text)))[:20]

        result["entities"] = {
            "emails": emails[:10],
            "urls": urls[:20],
            "potential_names": names,
            "dates": dates,
            "numbers": len(re.findall(r"\b\d+\b", text)),
        }

    return [TextContent(type="text", text=json.dumps(result, indent=2, default=str))]

def _count_syllables(word: str) -> int:
    """Approximate syllable count for English words."""
    word = word.lower().strip(".:;?!")
    if len(word) <= 3:
        return 1
    count = 0
    vowels = "aeiouy"
    prev_vowel = False
    for ch in word:
        is_vowel = ch in vowels
        if is_vowel and not prev_vowel:
            count += 1
        prev_vowel = is_vowel
    if word.endswith("e"):
        count -= 1
    if word.endswith("le") and len(word) > 2 and word[-3] not in vowels:
        count += 1
    return max(1, count)

@server.tool()
async def markdown_convert(
    content: str,
    from_format: str = "html",
    to_format: str = "plain",
) -> list[TextContent]:
    """Convert between markdown, HTML, and plain text.

    Args:
        content: Input content string
        from_format: "html", "markdown", or "plain"
        to_format: "html", "markdown", or "plain"
    """
    result = ""

    if from_format == "html" and to_format == "markdown":
        # Simple HTML-to-markdown
        text = content
        # Headers
        for i in range(6, 0, -1):
            text = re.sub(rf"<h{i}[^>]*>(.*?)</h{i}>", rf"{'#' * i} \1\n", text, flags=re.DOTALL | re.IGNORECASE)
        # Bold/Italic
        text = re.sub(r"<strong[^>]*>(.*?)</strong>", r"**\1**", text, flags=re.DOTALL | re.IGNORECASE)
        text = re.sub(r"<b[^>]*>(.*?)</b>", r"**\1**", text, flags=re.DOTALL | re.IGNORECASE)
        text = re.sub(r"<em[^>]*>(.*?)</em>", r"*\1*", text, flags=re.DOTALL | re.IGNORECASE)
        text = re.sub(r"<i[^>]*>(.*?)</i>", r"*\1*", text, flags=re.DOTALL | re.IGNORECASE)
        # Links
        text = re.sub(r'<a\s+href="([^"]+)"[^>]*>(.*?)</a>', r"[\2](\1)", text, flags=re.DOTALL | re.IGNORECASE)
        # Lists
        text = re.sub(r"<li[^>]*>(.*?)</li>", r"- \1\n", text, flags=re.DOTALL | re.IGNORECASE)
        # Code
        text = re.sub(r"<code[^>]*>(.*?)</code>", r"`\1`", text, flags=re.DOTALL | re.IGNORECASE)
        # Paragraphs
        text = re.sub(r"<p[^>]*>(.*?)</p>", r"\1\n\n", text, flags=re.DOTALL | re.IGNORECASE)
        # Strip remaining tags
        text = re.sub(r"<[^>]+>", "", text)
        result = text.strip()

    elif from_format == "markdown" and to_format == "html":
        # Basic markdown to HTML
        text = content
        text = re.sub(r"^### (.+)$", r"<h3>\1</h3>", text, flags=re.MULTILINE)
        text = re.sub(r"^## (.+)$", r"<h2>\1</h2>", text, flags=re.MULTILINE)
        text = re.sub(r"^# (.+)$", r"<h1>\1</h1>", text, flags=re.MULTILINE)
        text = re.sub(r"\*\*(.+?)\*\*", r"<strong>\1</strong>", text)
        text = re.sub(r"\*(.+?)\*", r"<em>\1</em>", text)
        text = re.sub(r"`(.+?)`", r"<code>\1</code>", text)
        text = re.sub(r"\[(.+?)\]\((.+?)\)", r'<a href="\2">\1</a>', text)
        text = re.sub(r"^- (.+)$", r"<li>\1</li>", text, flags=re.MULTILINE)
        lines = text.split("\n")
        result_lines = []
        in_list = False
        for line in lines:
            if line.startswith("<li>"):
                if not in_list:
                    result_lines.append("<ul>")
                    in_list = True
                result_lines.append(line)
            else:
                if in_list:
                    result_lines.append("</ul>")
                    in_list = False
                if line.strip():
                    result_lines.append(f"<p>{line}</p>")
                else:
                    result_lines.append("")
        if in_list:
            result_lines.append("</ul>")
        result = "\n".join(result_lines)

    elif to_format == "plain":
        # Strip all formatting
        result = re.sub(r"<[^>]+>", "", content)
        result = re.sub(r"[*_`#]", "", result)
        result = re.sub(r"\[(.+?)\]\(.+?\)", r"\1", result)

    else:
        result = content

    return [TextContent(type="text", text=result)]

@server.tool()
async def split_document(
    text: str,
    method: str = "paragraphs",
    chunk_size: int = 1000,
) -> list[TextContent]:
    """Split a long document into chunks for processing.

    Args:
        text: The document text
        method: "paragraphs", "sentences", "fixed" (character chunks), "markdown" (by headers)
        chunk_size: Target chunk size in characters (for "fixed" method)
    """
    chunks = []

    if method == "paragraphs":
        chunks = [p.strip() for p in text.split("\n\n") if p.strip()]
    elif method == "sentences":
        chunks = [s.strip() for s in re.split(r"[.!?]+\s+", text) if s.strip()]
    elif method == "fixed":
        for i in range(0, len(text), chunk_size):
            chunks.append(text[i:i + chunk_size])
    elif method == "markdown":
        # Split on headers
        parts = re.split(r"(?=^#{1,6}\s)", text, flags=re.MULTILINE)
        current = ""
        for part in parts:
            if len(current) + len(part) > chunk_size and current:
                chunks.append(current.strip())
                current = part
            else:
                current += part
        if current.strip():
            chunks.append(current.strip())

    return [TextContent(type="text", text=json.dumps({
        "method": method,
        "total_chunks": len(chunks),
        "chunk_sizes": [len(c) for c in chunks],
        "chunks": chunks,
    }, indent=2, ensure_ascii=False))]

@server.tool()
async def compare_documents(doc_a: str, doc_b: str) -> list[TextContent]:
    """Compare two documents and show similarity metrics.

    Args:
        doc_a: First document text
        doc_b: Second document text
    """
    import difflib

    words_a = set(re.findall(r"\b\w+\b", doc_a.lower()))
    words_b = set(re.findall(r"\b\w+\b", doc_b.lower()))

    intersection = words_a & words_b
    union = words_a | words_b
    jaccard = len(intersection) / len(union) if union else 0

    # Generate diff
    diff = list(difflib.unified_diff(
        doc_a.splitlines()[:100],
        doc_b.splitlines()[:100],
        lineterm="",
        fromfile="Document A",
        tofile="Document B",
    ))

    return [TextContent(type="text", text=json.dumps({
        "similarity": {
            "jaccard_coefficient": round(jaccard, 4),
            "shared_words": len(intersection),
            "doc_a_unique_words": len(words_a),
            "doc_b_unique_words": len(words_b),
            "only_in_a": len(words_a - words_b),
            "only_in_b": len(words_b - words_a),
        },
        "unified_diff": diff[:200],
    }, indent=2))]

# ---------------------------------------------------------------------------
# Resources
# ---------------------------------------------------------------------------

@server.resource("docs://formats")
async def supported_formats() -> str:
    return json.dumps({
        "read_formats": ["txt", "md", "json", "csv", "xml", "html", "pdf", "docx", "py", "js", "yaml", "toml"],
        "operations": ["parse", "extract_metadata", "text_analysis", "markdown_convert", "split_document", "compare_documents"],
    }, indent=2)

@server.resource("docs://document/{file_path}")
async def get_document_resource(file_path: str) -> str:
    """Parameterized resource — analyze any document by path and return its stats."""
    p = Path(file_path)
    if not p.exists():
        return json.dumps({"error": f"File not found: {file_path}"})
    text = p.read_text(errors="replace")
    words = text.split()
    return json.dumps({
        "file": file_path,
        "format": _detect_format(file_path),
        "size_bytes": p.stat().st_size,
        "lines": len(text.splitlines()),
        "words": len(words),
        "characters": len(text),
        "preview": text[:2000],
    }, indent=2, ensure_ascii=False)

# ---------------------------------------------------------------------------
# Prompts
# ---------------------------------------------------------------------------

@server.prompt(
    name="document-analyzer",
    description="Comprehensive document analysis and summarization workflow",
    arguments={
        "analysis_goal": {
            "type": "string",
            "enum": ["summarize", "compare", "extract_entities", "assess_readability", "full_audit"],
            "description": "What to do with the document",
            "required": True,
        },
    },
)
async def document_analyzer_prompt(analysis_goal: str) -> dict:
    workflows = {
        "summarize": "1) Parse the document. 2) Analyze text frequency to identify key terms. 3) Split into sections. 4) Extract key sentences from each section.",
        "compare": "1) Parse both documents. 2) Use compare_documents for similarity metrics. 3) Show the unified diff. 4) Highlight unique content in each.",
        "extract_entities": "1) Parse the document. 2) Use text_analysis with analysis_type='entities'. 3) Extract emails, URLs, names, dates, numbers. 4) Organize by entity type.",
        "assess_readability": "1) Parse the document. 2) Use text_analysis with analysis_type='readability'. 3) Report Flesch-Kincaid grade, reading ease, sentence stats. 4) Suggest simplifications if grade > 12.",
        "full_audit": "1) Parse and get metadata. 2) Full text analysis (frequency + readability + entities). 3) Check for issues. 4) Generate a comprehensive audit report.",
    }
    return {
        "messages": [{
            "role": "user",
            "content": {
                "type": "text",
                "text": f"""Analyze the provided document with goal: {analysis_goal}.

Workflow:
{workflows.get(analysis_goal, workflows['full_audit'])}

Use the document processing tools (parse_document, extract_metadata, text_analysis, compare_documents, split_document) to complete the analysis.

Present results as:
1. Executive summary (3 bullet points)
2. Detailed findings with metrics
3. Recommendations or action items""",
            },
        }],
    }

# ---------------------------------------------------------------------------
# Startup
# ---------------------------------------------------------------------------

async def main():
    transport = StdioServerTransport()
    await server.run(transport)

if __name__ == "__main__":
    asyncio.run(main())
